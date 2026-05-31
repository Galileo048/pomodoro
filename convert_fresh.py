# -*- coding: utf-8 -*-
"""
Create a FRESH document, copy content from original, convert formulas to OMath.
This avoids XML corruption from modifying the existing document.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

SRC = r'c:\Users\皮皮\Downloads\物理实验系统教程_带安装图.docx'
DST = r'c:\Users\皮皮\Downloads\物理实验系统教程_带安装图_OMath版2.docx'

OMATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NSMAP = {'m': OMATH_NS, 'w': W_NS}

_SUB_DIGITS = '0123456789'  # placeholder, will be overwritten
SUB_MAP = {}
for i, ch in enumerate('0123456789'):
    SUB_MAP[ch] = str(i)
# Actually use Unicode subscript chars
_SUB_UNI = '₀₁₂₃₄₅₆₇₈₉'
SUB_MAP = {}
for i, ch in enumerate(_SUB_UNI):
    SUB_MAP[ch] = str(i)
SUB_MAP['ᵢ'] = 'i'  # ᵢ
SUB_MAP['ₙ'] = 'n'  # ₙ

_SUP_UNI = '⁰¹²³⁴⁵⁶⁷⁸⁹'
SUP_MAP = {}
for i, ch in enumerate(_SUP_UNI):
    SUP_MAP[ch] = str(i)
SUP_MAP['⁻'] = '-'  # ⁻

FRAC_MAP = {
    '½': ('1', '2'),  # ½
    '⅓': ('1', '3'),  # ⅓
    '⅔': ('2', '3'),  # ⅔
    '¼': ('1', '4'),  # ¼
    '¾': ('3', '4'),  # ¾
}

ALL_SPECIAL = set(SUB_MAP) | set(SUP_MAP) | set(FRAC_MAP)


def has_special(text):
    return any(ch in ALL_SPECIAL for ch in text)


def _omath_r(text):
    r = etree.SubElement(etree.Element('dummy', nsmap=NSMAP), f'{{{OMATH_NS}}}r')
    rPr = etree.SubElement(r, f'{{{W_NS}}}rPr')
    rFonts = etree.SubElement(rPr, f'{{{W_NS}}}rFonts')
    rFonts.set(f'{{{W_NS}}}ascii', 'Cambria Math')
    rFonts.set(f'{{{W_NS}}}hAnsi', 'Cambria Math')
    t = etree.SubElement(r, f'{{{OMATH_NS}}}t')
    t.text = text
    t.set(f'{{{W_NS}}}space', 'preserve')
    return r


def make_omath_sub(base, sub_text):
    omath = etree.Element(f'{{{OMATH_NS}}}oMath', nsmap=NSMAP)
    sSub = etree.SubElement(omath, f'{{{OMATH_NS}}}sSub')
    e = etree.SubElement(sSub, f'{{{OMATH_NS}}}e'); e.append(_omath_r(base))
    sub = etree.SubElement(sSub, f'{{{OMATH_NS}}}sub'); sub.append(_omath_r(sub_text))
    return omath


def make_omath_sup(base, sup_text):
    omath = etree.Element(f'{{{OMATH_NS}}}oMath', nsmap=NSMAP)
    sSup = etree.SubElement(omath, f'{{{OMATH_NS}}}sSup')
    e = etree.SubElement(sSup, f'{{{OMATH_NS}}}e'); e.append(_omath_r(base))
    sup = etree.SubElement(sSup, f'{{{OMATH_NS}}}sup'); sup.append(_omath_r(sup_text))
    return omath


def make_omath_frac(num, den):
    omath = etree.Element(f'{{{OMATH_NS}}}oMath', nsmap=NSMAP)
    f = etree.SubElement(omath, f'{{{OMATH_NS}}}f')
    num_e = etree.SubElement(f, f'{{{OMATH_NS}}}num'); num_e.append(_omath_r(num))
    den_e = etree.SubElement(f, f'{{{OMATH_NS}}}den'); den_e.append(_omath_r(den))
    return omath


def parse_segments(text):
    segments = []
    buf = []
    i, n = 0, len(text)
    while i < n:
        ch = text[i]
        if ch in SUB_MAP:
            if buf:
                segments.append(('text', ''.join(buf)))
                buf = []
            sub_chars = [SUB_MAP[ch]]
            i += 1
            while i < n and text[i] in SUB_MAP:
                sub_chars.append(SUB_MAP[text[i]])
                i += 1
            if segments and segments[-1][0] == 'text' and segments[-1][1] and segments[-1][1][-1].isalpha():
                base = segments[-1][1][-1]
                segments[-1] = ('text', segments[-1][1][:-1])
                if not segments[-1][1]:
                    segments.pop()
                segments.append(('sub', base, ''.join(sub_chars)))
            else:
                segments.append(('sub_raw', ''.join(sub_chars)))
            continue
        elif ch in SUP_MAP:
            if buf:
                segments.append(('text', ''.join(buf)))
                buf = []
            sup_chars = [SUP_MAP[ch]]
            i += 1
            while i < n and text[i] in SUP_MAP:
                sup_chars.append(SUP_MAP[text[i]])
                i += 1
            if segments and segments[-1][0] == 'text' and segments[-1][1] and segments[-1][1][-1].isalpha():
                base = segments[-1][1][-1]
                segments[-1] = ('text', segments[-1][1][:-1])
                if not segments[-1][1]:
                    segments.pop()
                segments.append(('sup', base, ''.join(sup_chars)))
            else:
                segments.append(('sup_raw', ''.join(sup_chars)))
            continue
        elif ch in FRAC_MAP:
            if buf:
                segments.append(('text', ''.join(buf)))
                buf = []
            num, den = FRAC_MAP[ch]
            segments.append(('frac', num, den))
            i += 1
            continue
        else:
            buf.append(ch)
            i += 1
    if buf:
        segments.append(('text', ''.join(buf)))
    return segments


def convert_run_text(text, rPr_xml=None):
    """Convert a run's text to a list of w:r and m:oMath elements."""
    segments = parse_segments(text)
    if not segments:
        return None
    elems = []
    for seg in segments:
        if seg[0] == 'text':
            if seg[1]:
                r = OxmlElement('w:r')
                if rPr_xml:
                    rPr = etree.fromstring(rPr_xml)
                    r.append(rPr)
                t = OxmlElement('w:t')
                t.text = seg[1]
                t.set(qn('xml:space'), 'preserve')
                r.append(t)
                elems.append(r)
        elif seg[0] == 'sub':
            elems.append(make_omath_sub(seg[1], seg[2]))
        elif seg[0] == 'sup':
            elems.append(make_omath_sup(seg[1], seg[2]))
        elif seg[0] == 'sub_raw':
            elems.append(make_omath_sub(' ', seg[1]))
        elif seg[0] == 'sup_raw':
            elems.append(make_omath_sup(' ', seg[1]))
        elif seg[0] == 'frac':
            elems.append(make_omath_frac(seg[1], seg[2]))
    return elems


def copy_run_with_conversion(run_elem, parent):
    """Copy a run, converting any sub/sup/frac to OMath if needed."""
    t_elem = run_elem.find(qn('w:t'))
    if t_elem is None or not t_elem.text:
        # No text, just copy as-is
        parent.append(etree.fromstring(etree.tostring(run_elem)))
        return 0

    text = t_elem.text
    if not has_special(text):
        # No special chars, copy as-is
        parent.append(etree.fromstring(etree.tostring(run_elem)))
        return 0

    # Has special chars - convert
    rPr = run_elem.find(qn('w:rPr'))
    rPr_xml = etree.tostring(rPr, encoding='unicode') if rPr is not None else None

    # Merge with preceding text run if subscript at start
    elems = convert_run_text(text, rPr_xml)
    if elems:
        for elem in elems:
            parent.append(elem)
        return 1
    return 0


def process_paragraph_content(src_para, dst_para):
    """Copy paragraph content, converting formulas."""
    dst_elem = dst_para._element
    src_elem = src_para._element

    # Copy pPr (paragraph properties)
    pPr = src_elem.find(qn('w:pPr'))
    if pPr is not None:
        dst_elem.append(etree.fromstring(etree.tostring(pPr)))

    # Process runs and other elements
    runs = list(src_elem.findall(qn('w:r')))
    for run in runs:
        copy_run_with_conversion(run, dst_elem)


# ============================================================
print("[1/3] Loading original document...")
orig = Document(SRC)
print(f"  {len(orig.paragraphs)} paragraphs, {len(orig.tables)} tables")

print("[2/3] Creating fresh document with converted formulas...")
doc = Document()

# Copy styles from original
for style in orig.styles:
    try:
        if style.name not in [s.name for s in doc.styles]:
            doc.styles.add_style(style.name, style.type)
    except:
        pass

converted = 0

# Copy paragraphs
for i, src_para in enumerate(orig.paragraphs):
    # Create new paragraph with same style
    dst_para = doc.add_paragraph()
    try:
        dst_para.style = src_para.style
    except:
        pass

    # Copy paragraph content with formula conversion
    process_paragraph_content(src_para, dst_para)

    if i % 500 == 0:
        print(f"  Processed {i}/{len(orig.paragraphs)} paragraphs...")

# Count converted formulas
for p in doc.paragraphs:
    for child in p._element:
        if child.tag.split('}')[-1] == 'oMath':
            converted += 1

print(f"  OMath formulas in new doc: {converted}")

print("[3/3] Saving...")
doc.save(DST)
print(f"Saved to: {DST}")
