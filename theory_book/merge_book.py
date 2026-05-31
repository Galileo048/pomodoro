#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""合并所有章节为docx文件"""

import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    doc.add_page_break()

def create_cover(doc):
    """创建封面"""
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('看穿这个世界')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('马列毛与西马理论通俗读本')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run('━' * 30)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run('批判社会学与意识形态分析入门')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for _ in range(6):
        doc.add_paragraph()

    year = doc.add_paragraph()
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = year.add_run('二〇二五年')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    add_page_break(doc)

def main():
    # 章节文件列表
    chapters = [
        '01_引言.md',
        '02_第一章.md',
        '03_第二章.md',
        '04_第三章.md',
        '05_第四章.md',
        '06_第五章.md',
        '07_第六章.md',
        '08_第七章.md',
        '09_第八章.md',
        '10_第九章.md',
        '11_第十章.md',
        '12_第十一章.md',
        '13_第十二章.md',
        '14_第十三章.md',
        '15_第十四章.md',
        '16_第十五章.md',
        '17_第十六章.md',
        '18_第十七章.md',
        '19_第十八章.md',
        '20_第十九章.md',
        '21_第二十章.md',
        '22_第二十一章.md',
        '23_第二十二章.md',
        '24_第二十三章.md',
        '26_第二十四章.md',
        '27_第二十五章.md',
        '28_第二十六章.md',
        '29_第二十七章.md',
        '30_第二十八章.md',
        '31_第二十九章.md',
        '32_第三十章.md',
        '33_第三十一章.md',
        '34_第三十二章.md',
        '35_第三十三章.md',
        '36_第三十四章.md',
        '37_第三十五章.md',
        '38_第三十六章.md',
        '39_第三十七章.md',
        '40_第三十八章.md',
        '25_结语.md',  # 结语放在最后
    ]

    # 创建新文档
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # 设置页面
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    # 创建封面
    print("创建封面...")
    create_cover(doc)

    # 读取并合并所有章节
    base_dir = r'D:\pythonapp\open\theory_book'
    total_chars = 0
    chinese_chars = 0

    for chapter_file in chapters:
        filepath = os.path.join(base_dir, chapter_file)
        if not os.path.exists(filepath):
            print(f"警告: {chapter_file} 不存在，跳过")
            continue

        print(f"读取 {chapter_file}...")
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()

        total_chars += len(content)
        chinese_chars += sum(1 for c in content if '一' <= c <= '鿿')

        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 处理标题
            if line.startswith('# ') and not line.startswith('## '):
                heading = doc.add_heading(line[2:], level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line.strip('*'))
                run.bold = True
            elif line.startswith('---'):
                doc.add_paragraph('—' * 50)
            elif line.startswith('|'):
                doc.add_paragraph(line)
            elif line.startswith('**注释') or line.startswith('**参考文献'):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
            elif line.startswith('¹') or line.startswith('²') or line.startswith('³') or line.startswith('⁴') or line.startswith('⁵'):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.size = Pt(9)
            else:
                doc.add_paragraph(line)

    # 添加页眉页脚
    print("添加页眉页脚...")
    for section in doc.sections:
        # 页眉
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run('看穿这个世界——马列毛与西马理论通俗读本')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.italic = True

        # 页眉下划线
        pPr = hp._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '999999')
        pBdr.append(bottom)
        pPr.append(pBdr)

        # 页脚
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = fp.add_run('— ')
        run1.font.size = Pt(9)
        run1.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run1._r.addnext(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        fldChar1.addnext(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        instrText.addnext(fldChar2)

        run2 = fp.add_run(' —')
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # 保存
    output_path = os.path.join(base_dir, '看穿这个世界_完整版.docx')
    print(f"\n=== 统计 ===")
    print(f"总字符数: {total_chars}")
    print(f"中文字符数: {chinese_chars}")
    print(f"\n保存至: {output_path}")
    doc.save(output_path)
    print("完成！")

if __name__ == '__main__':
    main()
