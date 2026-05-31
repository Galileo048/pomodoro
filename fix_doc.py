# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SRC = r'c:\Users\皮皮\Desktop\四合一智能物理实验系统_完整教程_扩展版 (2).docx'
DST = r'c:\Users\皮皮\Desktop\四合一智能物理实验系统_完整教程_扩展版_修正版.docx'
DIAGRAMS = r'D:\pythonapp\open\diagrams_content.txt'

print("[1/6] Loading document...")
doc = Document(SRC)
print(f"  {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

# ============================================================
# FIX 1: Revert Table 1 to original correct data
# ============================================================
print("[2/6] Restoring Table 1 to correct data...")
t1 = doc.tables[1]
if t1.rows[0].cells[0].text.strip() == '实验名称':  # 实验名称
    correct_data = [
        ['平抛运动', '初速度、轨迹',
         '运动分解+图像追踪', '速度误差<3%'],
        ['简谐运动', '周期、弹簧常数',
         '光电门计时+正弦拟合', '周期误差<1%'],
        ['自由落体', '重力加速度g',
         '光电门计时+位移测量', 'g值误差<2%'],
        ['单摆', '重力加速度g',
         'T²-L关系+线性拟合', 'g值误差<2%'],
    ]
    for i, row_data in enumerate(correct_data):
        row_idx = i + 1
        if row_idx < len(t1.rows):
            for j, val in enumerate(row_data):
                cell = t1.rows[row_idx].cells[j]
                # Clear existing content
                for para in cell.paragraphs:
                    para.text = ""
                cell.paragraphs[0].text = val
    print("  Table 1 restored to correct data")
else:
    print("  WARNING: Table 1 header not found")

# ============================================================
# FIX 2: Typo physicis_lab_env -> physics_lab_env (P695)
# ============================================================
print("[3/6] Fixing typos...")
typo_count = 0
for para in doc.paragraphs:
    if 'physicis_lab_env' in para.text:
        for run in para.runs:
            if 'physicis_lab_env' in run.text:
                run.text = run.text.replace('physicis_lab_env', 'physics_lab_env')
                typo_count += 1
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if 'physicis_lab_env' in cell.text:
                cell.text = cell.text.replace('physicis_lab_env', 'physics_lab_env')
                typo_count += 1
print(f"  Fixed {typo_count} typo(s)")

# ============================================================
# FIX 3: Pendulum period wording (P607)
# ============================================================
print("[4/6] Fixing pendulum period wording...")
fix_count = 0
for para in doc.paragraphs:
    if '周期是从-θ₀到θ₀再返回的时间的四分之一的4倍' in para.text:
        # Text may be split across runs - rebuild the paragraph
        old_full = para.text
        # Find the old substring pattern
        if '周期是从-θ₀到θ₀再返回的时间的四分之一的4倍' in old_full:
            new_full = old_full.replace(
                '周期是从-θ₀到θ₀再返回的时间的四分之一的4倍',
                '周期是从θ=0运动到θ=θ₀所需时间的4倍（利用摆动的对称性）'
            )
            # Clear all runs and set new text in first run
            if para.runs:
                para.runs[0].text = new_full
                for r in para.runs[1:]:
                    r.text = ''
            fix_count += 1
            print("  Fixed pendulum period wording")
            break

# ============================================================
# FIX 4: Over-damped code initial condition note (P512)
# ============================================================
print("[5/6] Adding initial condition note to over-damped code...")
ic_count = 0
for para in doc.paragraphs:
    if 'A = X0 / 2' in para.text and 'B = X0 / 2' in para.text:
        for run in para.runs:
            if 'B = X0 / 2' in run.text and 'x(0)' not in run.text:
                run.text = run.text.replace(
                    'B = X0 / 2',
                    'B = X0 / 2  # 注：此解对应初始条件 x(0)=X₀, v(0)=0'
                )
                ic_count += 1
                print("  Added initial condition note")
                break
        if ic_count == 0:
            # Try alternate formatting
            for run in para.runs:
                if 'B = X0/2' in run.text and 'x(0)' not in run.text:
                    run.text = run.text.replace(
                        'B = X0/2',
                        'B = X0/2  # 注：此解对应初始条件 x(0)=X₀, v(0)=0'
                    )
                    ic_count += 1
                    print("  Added initial condition note (alt format)")
                    break

print(f"  Text fixes total: {fix_count + ic_count + typo_count}")

# ============================================================
# PART 2: Insert diagram sections at correct positions
# ============================================================
print("[6/6] Inserting diagram sections...")

# Read diagrams content
with open(DIAGRAMS, 'r', encoding='utf-8') as f:
    diagrams_text = f.read()

# Parse sections
sections = []
current_section = None
current_title = None
current_paras = []

for line in diagrams_text.split('\n'):
    line = line.rstrip()
    if line.startswith('==SECTION:'):
        if current_section:
            sections.append({
                'id': current_section,
                'title': current_title,
                'paras': current_paras
            })
        current_section = line.replace('==SECTION:', '').replace('==', '')
        current_title = None
        current_paras = []
    elif line.startswith('==TITLE:'):
        current_title = line.replace('==TITLE:', '').replace('==', '')
    elif line.startswith('==PARA:'):
        current_paras.append(('text', line.replace('==PARA:', '').replace('==', '')))
    elif line.strip() and current_section:
        current_paras.append(('diagram', line))

if current_section:
    sections.append({
        'id': current_section,
        'title': current_title,
        'paras': current_paras
    })

print(f"  Parsed {len(sections)} diagram sections")

# Define insertion points by searching for heading text (skip TOC entries)
insertion_points = {}
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name
    if 'toc' in style.lower():
        continue

    # Framework assembly - P1105
    if '28.2' in text and '框架组装' in text and 'Heading' in style:
        insertion_points['FRAME_ASSEMBLY'] = i
    # Release mechanism - P1159
    if '28.3' in text and '发射装置' in text and 'Heading' in style:
        insertion_points['RELEASE_MECHANISM'] = i
    # Photogate - P1574
    if '29.4' in text and '光电门' in text and 'Heading' in style:
        insertion_points['PHOTOGATE_INSTALL'] = i
    # Circuit wiring - P1718 (30.1 电磁铁释放装置详解)
    if '30.1' in text and '电磁铁' in text and 'Heading' in style:
        insertion_points['CIRCUIT_WIRING'] = i
    # ESP32 GPIO - P848 (11.3)
    if '11.3' in text and '引脚' in text and 'Heading' in style:
        insertion_points['ESP32_WIRING_SUMMARY'] = i

print(f"  Found insertion points: {list(insertion_points.keys())}")

# Helper: create a new paragraph element with heading style
def make_heading_para(text, level=3):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), f'Heading {level}')
    pPr.append(pStyle)
    p.append(pPr)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rBold = OxmlElement('w:b')
    rPr.append(rBold)
    rColor = OxmlElement('w:color')
    rColor.set(qn('w:val'), '2E75B6')
    rPr.append(rColor)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    p.append(run)
    return p

# Helper: create a diagram paragraph (monospace)
def make_diagram_para(text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '20')
    spacing.set(qn('w:after'), '20')
    pPr.append(spacing)
    p.append(pPr)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFont = OxmlElement('w:rFonts')
    rFont.set(qn('w:ascii'), 'Consolas')
    rFont.set(qn('w:hAnsi'), 'Consolas')
    rFont.set(qn('w:eastAsia'), '等线')
    rPr.append(rFont)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '18')
    rPr.append(sz)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    p.append(run)
    return p

# Helper: create a normal text paragraph
def make_text_para(text):
    p = OxmlElement('w:p')
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    p.append(run)
    return p

# Insert each section
total_inserted = 0
for section in sections:
    target_idx = insertion_points.get(section['id'])
    if target_idx is None:
        print(f"  WARNING: No insertion point for '{section['id']}', appending at end")
        target_para = None
    else:
        target_para = doc.paragraphs[target_idx]

    # Build list of XML elements to insert
    elements = []
    # Title
    elements.append(make_heading_para(section['title']))
    # Content
    for ptype, content in section['paras']:
        if ptype == 'diagram':
            elements.append(make_diagram_para(content))
        else:
            elements.append(make_text_para(content))

    # Insert after target
    if target_para is not None:
        last_elem = target_para._element
        for elem in elements:
            last_elem.addnext(elem)
            last_elem = elem
            total_inserted += 1
        print(f"  Inserted '{section['id']}' ({len(elements)} paras) after P{target_idx}")
    else:
        # Append at end
        for elem in elements:
            doc.element.body.append(elem)
            total_inserted += 1
        print(f"  Appended '{section['id']}' ({len(elements)} paras) at end")

print(f"  Total paragraphs inserted: {total_inserted}")

# ============================================================
# Save
# ============================================================
print("Saving document...")
doc.save(DST)
print(f"Saved to: {DST}")
print("Done!")
