# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
import re

src = r'c:\Users\皮皮\Downloads\Manim动画分镜脚本_第1-5期.docx'

doc = Document(src)

# Extract ALL paragraphs with styles
print("=" * 80)
print("FULL DOCUMENT CONTENT")
print("=" * 80)

for i, para in enumerate(doc.paragraphs):
    try:
        style = para.style.name if para.style else "None"
    except:
        style = "Unknown"
    text = para.text.strip()
    if text:
        print(f"\n[P{i:03d}] [{style}] {text}")

# Extract tables
print("\n" + "=" * 80)
print("TABLES")
print("=" * 80)

for ti, table in enumerate(doc.tables):
    print(f"\n--- Table {ti} ---")
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  Row {ri}: {cells}")

# Extract images info
print("\n" + "=" * 80)
print("IMAGES INFO")
print("=" * 80)

image_count = 0
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        image_count += 1
        print(f"  Image {image_count}: {rel.target_ref}")

print(f"\nTotal images: {image_count}")

# Check for inline shapes
inline_count = 0
for para in doc.paragraphs:
    for run in para.runs:
        if run._element.findall('.//' + qn('w:drawing')):
            inline_count += 1
print(f"Total inline drawings: {inline_count}")