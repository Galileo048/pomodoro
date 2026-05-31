# -*- coding: utf-8 -*-
"""
Document processor using OMath (Word equation editor) for subscripts/superscripts/fractions.
1. Apply error fixes
2. Insert diagram sections
3. Convert Unicode sub/sup/frac chars to Word OMath elements
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree
import re

SRC = r'c:\Users\皮皮\Desktop\四合一智能物理实验系统_完整教程_扩展版 (2).docx'
DST = r'c:\Users\皮皮\Desktop\四合一智能物理实验系统_完整教程_扩展版_修正版.docx'
DIAGRAMS = r'D:\pythonapp\open\diagrams_content.txt'

OMATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NSMAP = {'m': OMATH_NS, 'w': W_NS}

# Character maps
SUB_MAP = {}
for ch in '₀₁₂₃₄₅₆₇₈₉':
    SUB_MAP[ch] = str('₀₁₂₃₄₅₆₇₈₉'.index(ch))
SUB_MAP['ᵢ'] = 'i'
SUB_MAP['ₙ'] = 'n'

SUP_MAP = {}
for ch in '⁰¹²³⁴⁵⁶⁷⁸⁹':
    SUP_MAP[ch] = str('⁰¹²³⁴⁵⁶⁷⁸⁹'.index(ch))
SUP_MAP['⁻'] = '-'

FRAC_MAP = {
    '½': ('1', '2'), '⅓': ('1', '3'), '⅔': ('2', '3'),
    '¼': ('1', '4'), '¾': ('3', '4'),
}

ALL_SPECIAL = set(SUB_MAP) | set(SUP_MAP) | set(FRAC_MAP)


def has_special(text):
    return any(ch in ALL_SPECIAL for ch in text)


def merge_split_subscripts(para_elem):
    """
    Pre-process: when a run starts with sub/sup char and preceding run ends with alpha,
    move the alpha char to the start of the current run so subscript has a base.
    """
    runs = list(para_elem.findall(qn('w:r')))
    i = 1
    while i < len(runs):
        curr_t = runs[i].find(qn('w:t'))
        prev_t = runs[i-1].find(qn('w:t'))
        if curr_t is not None and curr_t.text and prev_t is not None and prev_t.text:
            curr_text = curr_t.text
            prev_text = prev_t.text
            if curr_text and curr_text[0] in ALL_SPECIAL and prev_text and prev_text[-1].isalpha():
                # Move last alpha char from prev to start of curr
                prev_t.text = prev_text[:-1]
                curr_t.text = prev_text[-1] + curr_text
        i += 1


def copy_rPr(src_run_elem):
    """Copy run properties from an existing run element."""
    rPr = src_run_elem.find(qn('w:rPr'))
    if rPr is not None:
        return etree.fromstring(etree.tostring(rPr))
    return None


def make_text_run(text, rPr_copy=None):
    """Create a plain w:r text element."""
    r = OxmlElement('w:r')
    if rPr_copy is not None:
        r.append(rPr_copy)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    return r


def _make_omath_r(text):
    """Create an m:r element with m:t inside, with Cambria Math font."""
    r = etree.SubElement(etree.Element('dummy', nsmap=NSMAP), f'{{{OMATH_NS}}}r')
    rPr = etree.SubElement(r, f'{{{W_NS}}}rPr')
    rFonts = etree.SubElement(rPr, f'{{{W_NS}}}rFonts')
    rFonts.set(f'{{{W_NS}}}ascii', 'Cambria Math')
    rFonts.set(f'{{{W_NS}}}hAnsi', 'Cambria Math')
    t = etree.SubElement(r, f'{{{OMATH_NS}}}t')
    t.text = text
    t.set(f'{{{W_NS}}}space', 'preserve')
    return r


def make_omath_sub(base, sub_text):
    """Create m:oMath for subscript: base_{sub}"""
    omath = etree.Element(f'{{{OMATH_NS}}}oMath', nsmap=NSMAP)
    sSub = etree.SubElement(omath, f'{{{OMATH_NS}}}sSub')
    e = etree.SubElement(sSub, f'{{{OMATH_NS}}}e')
    e.append(_make_omath_r(base))
    sub = etree.SubElement(sSub, f'{{{OMATH_NS}}}sub')
    sub.append(_make_omath_r(sub_text))
    return omath


def make_omath_sup(base, sup_text):
    """Create m:oMath for superscript: base^{sup}"""
    omath = etree.Element(f'{{{OMATH_NS}}}oMath', nsmap=NSMAP)
    sSup = etree.SubElement(omath, f'{{{OMATH_NS}}}sSup')
    e = etree.SubElement(sSup, f'{{{OMATH_NS}}}e')
    e.append(_make_omath_r(base))
    sup = etree.SubElement(sSup, f'{{{OMATH_NS}}}sup')
    sup.append(_make_omath_r(sup_text))
    return omath


def make_omath_frac(num, den):
    """Create m:oMath for fraction: num/den"""
    omath = etree.Element(f'{{{OMATH_NS}}}oMath', nsmap=NSMAP)
    f = etree.SubElement(omath, f'{{{OMATH_NS}}}f')
    num_e = etree.SubElement(f, f'{{{OMATH_NS}}}num')
    num_e.append(_make_omath_r(num))
    den_e = etree.SubElement(f, f'{{{OMATH_NS}}}den')
    den_e.append(_make_omath_r(den))
    return omath


def make_omath_combined(segments):
    """
    Create a single m:oMath element from multiple segments.
    segments: list of ('text', str), ('sub', base, sub), ('sup', base, sup), ('frac', num, den)
    """
    omath = etree.Element(f'{{{OMATH_NS}}}oMath', nsmap=NSMAP)
    for seg in segments:
        if seg[0] == 'text':
            omath.append(_make_omath_r(seg[1]))
        elif seg[0] == 'sub':
            sSub = etree.SubElement(omath, f'{{{OMATH_NS}}}sSub')
            e = etree.SubElement(sSub, f'{{{OMATH_NS}}}e')
            e.append(_make_omath_r(seg[1]))
            sub = etree.SubElement(sSub, f'{{{OMATH_NS}}}sub')
            sub.append(_make_omath_r(seg[2]))
        elif seg[0] == 'sup':
            sSup = etree.SubElement(omath, f'{{{OMATH_NS}}}sSup')
            e = etree.SubElement(sSup, f'{{{OMATH_NS}}}e')
            e.append(_make_omath_r(seg[1]))
            sup = etree.SubElement(sSup, f'{{{OMATH_NS}}}sup')
            sup.append(_make_omath_r(seg[2]))
        elif seg[0] == 'frac':
            f_elem = etree.SubElement(omath, f'{{{OMATH_NS}}}f')
            num_e = etree.SubElement(f_elem, f'{{{OMATH_NS}}}num')
            num_e.append(_make_omath_r(seg[1]))
            den_e = etree.SubElement(f_elem, f'{{{OMATH_NS}}}den')
            den_e.append(_make_omath_r(seg[2]))
    return omath


def parse_text_segments(text):
    """Parse text into segments: ('text',s), ('sub',base,sub), ('sup',base,sup), ('frac',num,den)."""
    segments = []
    buf = []
    i = 0
    n = len(text)
    while i < n:
        ch = text[i]
        if ch in SUB_MAP:
            if buf:
                segments.append(('text', ''.join(buf)))
                buf = []
            sub_chars = [SUB_MAP[ch]]
            i += 1
            while i < n and text[i] in SUB_MAP:
                sub_chars.append(SUB_MAP[text[i]])
                i += 1
            # Try to grab preceding alpha char as base
            if segments and segments[-1][0] == 'text' and segments[-1][1] and segments[-1][1][-1].isalpha():
                base_char = segments[-1][1][-1]
                segments[-1] = ('text', segments[-1][1][:-1])
                if not segments[-1][1]:
                    segments.pop()
                segments.append(('sub', base_char, ''.join(sub_chars)))
            else:
                segments.append(('sub_raw', ''.join(sub_chars)))
            continue
        elif ch in SUP_MAP:
            if buf:
                segments.append(('text', ''.join(buf)))
                buf = []
            sup_chars = [SUP_MAP[ch]]
            i += 1
            while i < n and text[i] in SUP_MAP:
                sup_chars.append(SUP_MAP[text[i]])
                i += 1
            if segments and segments[-1][0] == 'text' and segments[-1][1] and segments[-1][1][-1].isalpha():
                base_char = segments[-1][1][-1]
                segments[-1] = ('text', segments[-1][1][:-1])
                if not segments[-1][1]:
                    segments.pop()
                segments.append(('sup', base_char, ''.join(sup_chars)))
            else:
                segments.append(('sup_raw', ''.join(sup_chars)))
            continue
        elif ch in FRAC_MAP:
            if buf:
                segments.append(('text', ''.join(buf)))
                buf = []
            num, den = FRAC_MAP[ch]
            segments.append(('frac', num, den))
            i += 1
            continue
        else:
            buf.append(ch)
            i += 1
    if buf:
        segments.append(('text', ''.join(buf)))
    return segments


def process_run(run_elem, parent):
    """Process a single run: replace Unicode sub/sup/frac with OMath elements."""
    text_elem = run_elem.find(qn('w:t'))
    if text_elem is None or not text_elem.text:
        return 0
    text = text_elem.text
    if not has_special(text):
        return 0

    rPr = copy_rPr(run_elem)
    segments = parse_text_segments(text)
    if not segments:
        return 0

    # Build new elements: text runs + OMath blocks
    run_index = list(parent).index(run_elem)
    new_elems = []
    formula_count = 0

    for seg in segments:
        if seg[0] == 'text':
            if seg[1]:
                new_elems.append(make_text_run(seg[1], etree.fromstring(etree.tostring(rPr)) if rPr is not None else None))
        elif seg[0] == 'sub':
            # OMath sSub already contains the base char, no need for extra text run
            new_elems.append(make_omath_sub(seg[1], seg[2]))
            formula_count += 1
        elif seg[0] == 'sup':
            new_elems.append(make_omath_sup(seg[1], seg[2]))
            formula_count += 1
        elif seg[0] == 'sub_raw':
            new_elems.append(make_omath_sub(' ', seg[1]))
            formula_count += 1
        elif seg[0] == 'sup_raw':
            new_elems.append(make_omath_sup(' ', seg[1]))
            formula_count += 1
        elif seg[0] == 'frac':
            new_elems.append(make_omath_frac(seg[1], seg[2]))
            formula_count += 1

    # Remove original run
    parent.remove(run_elem)
    # Insert new elements
    for j, elem in enumerate(new_elems):
        parent.insert(run_index + j, elem)

    return formula_count


# ============================================================
# OXml paragraph builders for diagrams
# ============================================================
def make_heading_para(text, level=3):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), f'Heading {level}')
    pPr.append(pStyle); p.append(pPr)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rBold = OxmlElement('w:b'); rPr.append(rBold)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text; t.set(qn('xml:space'), 'preserve')
    run.append(t); p.append(run)
    return p

def make_diagram_para(text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '20'); spacing.set(qn('w:after'), '20')
    pPr.append(spacing); p.append(pPr)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFont = OxmlElement('w:rFonts')
    rFont.set(qn('w:ascii'), 'Consolas'); rFont.set(qn('w:hAnsi'), 'Consolas')
    rFont.set(qn('w:eastAsia'), '等线'); rPr.append(rFont)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '18'); rPr.append(sz)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text; t.set(qn('xml:space'), 'preserve')
    run.append(t); p.append(run)
    return p

def make_text_para(text):
    p = OxmlElement('w:p')
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr'); run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text; t.set(qn('xml:space'), 'preserve')
    run.append(t); p.append(run)
    return p


# ============================================================
# Main
# ============================================================
print("[1/4] Loading document...")
doc = Document(SRC)
print(f"  {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

# FIX 1: Table 1
print("[2/4] Applying error fixes...")
t1 = doc.tables[1]
if t1.rows[0].cells[0].text.strip() == '实验名称':
    correct_data = [
        ['平抛运动', '初速度、轨迹', '运动分解+图像追踪', '速度误差<3%'],
        ['简谐运动', '周期、弹簧常数', '光电门计时+正弦拟合', '周期误差<1%'],
        ['自由落体', '重力加速度g', '光电门计时+位移测量', 'g值误差<2%'],
        ['单摆', '重力加速度g', 'T²-L关系+线性拟合', 'g值误差<2%'],
    ]
    for i, row_data in enumerate(correct_data):
        for j, val in enumerate(row_data):
            cell = t1.rows[i+1].cells[j]
            for para in cell.paragraphs:
                para.text = ""
            cell.paragraphs[0].text = val
    print("  Table 1 restored")

# FIX 2: Typo
for para in doc.paragraphs:
    if 'physicis_lab_env' in para.text:
        for run in para.runs:
            if 'physicis_lab_env' in run.text:
                run.text = run.text.replace('physicis_lab_env', 'physics_lab_env')
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if 'physicis_lab_env' in cell.text:
                cell.text = cell.text.replace('physicis_lab_env', 'physics_lab_env')
print("  Typo fixed")

# FIX 3: Pendulum wording
for para in doc.paragraphs:
    if '周期是从-θ₀到θ₀再返回的时间的四分之一的4倍' in para.text:
        old = para.text
        new = old.replace(
            '周期是从-θ₀到θ₀再返回的时间的四分之一的4倍',
            '周期是从θ=0运动到θ=θ₀所需时间的4倍（利用摆动的对称性）'
        )
        if para.runs:
            para.runs[0].text = new
            for r in para.runs[1:]:
                r.text = ''
        break
print("  Pendulum wording fixed")

# FIX 4: Over-damped initial condition
for para in doc.paragraphs:
    if 'A = X0 / 2' in para.text and 'B = X0 / 2' in para.text:
        for run in para.runs:
            if 'B = X0 / 2' in run.text and 'x(0)' not in run.text:
                run.text = run.text.replace(
                    'B = X0 / 2',
                    'B = X0 / 2  # 注：此解对应初始条件 x(0)=X₀, v(0)=0'
                )
                break
print("  Initial condition note added")

# Insert diagram sections
print("[3/4] Inserting diagram sections...")
with open(DIAGRAMS, 'r', encoding='utf-8') as f:
    diagrams_text = f.read()

sections = []
current_section = None; current_title = None; current_paras = []
for line in diagrams_text.split('\n'):
    line = line.rstrip()
    if line.startswith('==SECTION:'):
        if current_section:
            sections.append({'id': current_section, 'title': current_title, 'paras': current_paras})
        current_section = line.replace('==SECTION:', '').replace('==', '')
        current_title = None; current_paras = []
    elif line.startswith('==TITLE:'):
        current_title = line.replace('==TITLE:', '').replace('==', '')
    elif line.startswith('==PARA:'):
        current_paras.append(('text', line.replace('==PARA:', '').replace('==', '')))
    elif line.strip() and current_section:
        current_paras.append(('diagram', line))
if current_section:
    sections.append({'id': current_section, 'title': current_title, 'paras': current_paras})

insertion_points = {}
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip(); style = para.style.name
    if 'toc' in style.lower(): continue
    if '28.2' in text and '框架组装' in text and 'Heading' in style:
        insertion_points['FRAME_ASSEMBLY'] = i
    if '28.3' in text and '发射装置' in text and 'Heading' in style:
        insertion_points['RELEASE_MECHANISM'] = i
    if '29.4' in text and '光电门' in text and 'Heading' in style:
        insertion_points['PHOTOGATE_INSTALL'] = i
    if '30.1' in text and '电磁铁' in text and 'Heading' in style:
        insertion_points['CIRCUIT_WIRING'] = i
    if '11.3' in text and '引脚' in text and 'Heading' in style:
        insertion_points['ESP32_WIRING_SUMMARY'] = i

for section in sections:
    target_idx = insertion_points.get(section['id'])
    elements = [make_heading_para(section['title'])]
    for ptype, content in section['paras']:
        if ptype == 'diagram':
            elements.append(make_diagram_para(content))
        else:
            elements.append(make_text_para(content))
    if target_idx is not None:
        last_elem = doc.paragraphs[target_idx]._element
        for elem in elements:
            last_elem.addnext(elem); last_elem = elem
        print(f"  Inserted '{section['id']}' after P{target_idx}")
    else:
        for elem in elements:
            doc.element.body.append(elem)
        print(f"  Appended '{section['id']}' at end")

# Convert subscripts/superscripts/fractions to OMath
print("[4/4] Converting sub/sup/frac to OMath...")
converted = 0

# Process paragraphs
for i, p in enumerate(doc.paragraphs):
    style = p.style.name
    if 'toc' in style.lower():
        continue
    # Merge base chars split across runs
    merge_split_subscripts(p._element)
    runs = list(p._element.findall(qn('w:r')))
    for run_elem in runs:
        text_elem = run_elem.find(qn('w:t'))
        if text_elem is not None and text_elem.text and has_special(text_elem.text):
            n = process_run(run_elem, p._element)
            converted += n

# Process tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                merge_split_subscripts(para._element)
                runs = list(para._element.findall(qn('w:r')))
                for run_elem in runs:
                    text_elem = run_elem.find(qn('w:t'))
                    if text_elem is not None and text_elem.text and has_special(text_elem.text):
                        n = process_run(run_elem, para._element)
                        converted += n

print(f"  OMath formulas created: {converted}")

print("Saving document...")
doc.save(DST)
print(f"Saved to: {DST}")
print("Done!")
