import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ============ 标题页 ============
title = doc.add_heading('基于Manim与AI辅助的中学数理可视化教学项目', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('项目规划 · 学习路径 · 时间安排 · 注意事项')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('整理自 DeepSeek 对话记录\n2026年5月30日')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ============ 目录 ============
doc.add_heading('目录', level=1)
toc_items = [
    '一、项目总体概述',
    '二、项目技术路线与核心架构',
    '三、详细时间安排（暑假 + 大三上 + 大三下）',
    '四、学习路径规划（Python + Manim + HTML/CSS + Uniapp）',
    '五、团队分工与协作方案',
    '六、关键注意事项与避坑指南',
    '七、DeepSeek API 使用策略',
    '八、大创申报策略与技巧',
    '九、每日作息与学习节奏建议',
    '十、阶段性检查清单',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ============ 一、项目总体概述 ============
doc.add_heading('一、项目总体概述', level=1)

doc.add_heading('1.1 项目名称（建议）', level=2)
doc.add_paragraph('基于Manim与AI辅助的中学生科学可视化互动学习APP')

doc.add_heading('1.2 项目简介', level=2)
doc.add_paragraph(
    '本项目针对中学物理和数学中大量抽象概念（如函数图像变换、物体运动分解、几何关系推导）'
    '学生难以直观理解的痛点，旨在开发一套高质量、可视化的动画教学视频资源。'
    '项目采用开源数学动画引擎Manim（Community Edition）作为核心制作工具，'
    '并创新性地引入DeepSeek大语言模型API辅助生成与调试动画代码，大幅降低师范生进行程序化动画创作的技术门槛。'
)

doc.add_heading('1.3 核心成果目标', level=2)
results = [
    '10-20个Manim动画教学视频（每集1-3分钟），投放到B站形成公开合集',
    '个人教学资源网站1个，实现视频分类展示、推导说明与课件下载',
    '可选：跨平台学习APP 1个，支持参数化交互（滑块操控动画变量）',
    'GitHub开源代码仓库，展示完整开发过程',
    '项目研究报告1份，争取发表教育技术相关论文1篇',
]
for r in results:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('1.4 项目创新点', level=2)
innovations = [
    '技术融合创新：将Manim程序化动画与DeepSeek/Claude大模型辅助编程结合，构建"自然语言描述→AI生成代码→人工审校调参"的新型教育资源生产流程',
    '双载体传播模式：以B站作为轻量化公开传播主阵地，以个人静态网站作为深度内容沉淀和课件归档平台',
    '交互模式创新（APP）：突破传统动画视频的"单向灌输"模式，引入变量操控机制，将抽象概念转化为可触摸、可探索的动态视觉对象',
    '学科交叉落地：物理、数学师范生主导教学内容，计算机、网工专业队员提供技术支持',
]
for inn in innovations:
    doc.add_paragraph(inn, style='List Bullet')

doc.add_heading('1.5 可行性分析', level=2)
doc.add_paragraph('【优势】', style='List Bullet')
advantages = [
    '数理可视化是刚需，大学都在推"数字化教学资源建设"',
    'Manim开源、社区版ManimCE稳定、教程多，上手成本逐年降低',
    '录像+配音即可，不需要出镜，制作门槛低',
    'B站播放量、弹幕、评论天然就是大创成果的"社会影响力"佐证',
    'AI辅助大幅降低编程门槛，师范生也能产出高质量动画',
]
for a in advantages:
    doc.add_paragraph(a, style='List Bullet 2')

doc.add_paragraph('【核心挑战与对策】', style='List Bullet')
challenges = [
    'Manim学习曲线陡 → 从3Blue1Brown开源项目找模板，先改参数再写原创',
    '编程基础薄弱 → 暑假6-8周集中学习，每天3小时即可产出第一个视频',
    '队员力量未充分发挥 → 精准分工，让每个人在最适合的位置发力',
    '大创评审风险 → 向APP方向包装升级，从"教育实践"变为"教育技术创新"',
]
for c in challenges:
    doc.add_paragraph(c, style='List Bullet 2')

doc.add_page_break()

# ============ 二、项目技术路线 ============
doc.add_heading('二、项目技术路线与核心架构', level=1)

doc.add_heading('2.1 总体技术路线', level=2)
flow_steps = [
    '阶段一：基础与工具链搭建 → Python + Manim环境 + DeepSeek API接入 + Git/GitHub初始化',
    '阶段二：高质量动画核心开发 → 选题 → 分镜脚本 → AI辅助代码生成 → 调试调优 → 渲染4K片段',
    '阶段三：后期合成与批量生产 → 配音录制 → 剪映精剪字幕 → 统一封面 → 质量审核',
    '阶段四：双载体发布 → B站上传合集 → 个人网站嵌入与课件归档 → GitHub开源',
    '阶段五：APP开发（可选） → Uniapp环境 → 参数化动画组件 → 本地视频打包与交互',
]
for s in flow_steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('2.2 核心工具链', level=2)
table = doc.add_table(rows=10, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['用途', '工具', '说明']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ['Python开发', 'PyCharm / Cursor', '代码编写、调试、AI补全'],
    ['动画制作', 'ManimCE 0.18+', '程序化数学动画引擎'],
    ['AI代码生成', 'DeepSeek API + Claude Code', '自然语言→Manim代码'],
    ['AI对话调试', 'Cherry Studio / Cursor', '图形化API客户端'],
    ['前端开发', 'Trae / VS Code', 'HTML/CSS/JS编辑'],
    ['视频剪辑', '剪映', '配音对齐、字幕、封面'],
    ['版本控制', 'Git + GitHub', '代码备份、开源展示'],
    ['网站部署', 'GitHub Pages', '免费静态网站托管'],
    ['APP开发（可选）', 'Uniapp + HBuilder X', '跨平台APP开发'],
]
for i, row_data in enumerate(data):
    for j, cell_data in enumerate(row_data):
        table.rows[i+1].cells[j].text = cell_data

doc.add_heading('2.3 标准化视频生产流水线', level=2)
pipeline = [
    '选知识点 + 写脚本（你和A，1天）→ 从人教版教材目录里勾选一个核心概念，写出300字画外音，分5-6个镜头',
    '生成动画代码（你，0.5-1天）→ 将分镜脚本喂给DeepSeek API，获取初稿，本地运行调试改参数',
    '渲染与录音（你，0.5天）→ 用 -qh 渲染高质量视频，用手机或电脑麦克风对照时间轴录音',
    '合成与剪辑（你，1h）→ 在剪映中对齐音频与动画，添加背景音乐和字幕',
    '发布与归档（你+C，0.5天）→ 上传B站填写标签简介，网站添加新视频的iframe和说明文字',
]
for p_text in pipeline:
    doc.add_paragraph(p_text, style='List Number')

doc.add_paragraph('按照这个流水线，后期每周可以稳定产出2个视频。')

doc.add_page_break()

# ============ 三、详细时间安排 ============
doc.add_heading('三、详细时间安排', level=1)

doc.add_heading('3.1 暑假总览（2026年7月-8月，共8周）', level=2)

summer_table = doc.add_table(rows=9, cols=3)
summer_table.style = 'Light Grid Accent 1'
summer_table.alignment = WD_TABLE_ALIGNMENT.CENTER

summer_headers = ['时间', '阶段', '核心目标']
for i, h in enumerate(summer_headers):
    summer_table.rows[0].cells[i].text = h

summer_data = [
    ['暑假前1周', '环境准备', 'Python+Manim+DeepSeek API环境全部跑通'],
    ['第1-2周', 'Python基础速通', '变量/循环/判断/函数/类/numpy基础，搭建开发环境'],
    ['第3-4周', 'Manim入门+第一个作品', 'Scene/Mobject/动画流程，产出第一个30秒运动学动画'],
    ['第5周', 'DeepSeek API融入工作流', '配置API，用标准Prompt生成Manim代码，建立AI工作流'],
    ['第6-7周', '批量产出（6-8个视频）', '脚本→AI生成→调试→渲染→配音→合成，每个视频2-3天'],
    ['第7周穿插', 'HTML/CSS速成（6h）', 'HTML骨架+CSS核心+B站iframe嵌入+响应式初步'],
    ['第8周', '网站搭建+项目收尾', 'GitHub Pages网站上线，代码开源准备，大创申报书草稿'],
    ['第8周+追加', 'Uniapp基础（可选）', 'Vue基础+Uniapp页面+第一个参数化交互Demo'],
]
for i, row_data in enumerate(summer_data):
    for j, cell_data in enumerate(row_data):
        summer_table.rows[i+1].cells[j].text = cell_data

doc.add_heading('3.2 暑假结束检查点', level=2)
checkpoints = [
    '✅ 有6-8个已上传B站的教学视频',
    '✅ 个人网站可以访问，每个视频都有独立介绍页面和嵌入播放器',
    '✅ GitHub仓库有完整的代码和课件，提交记录丰富',
    '✅ DeepSeek API已使用，能熟练用标准Prompt生成Manim代码',
    '✅ 队员A已提供至少10个视频的分镜脚本',
]
for c in checkpoints:
    doc.add_paragraph(c)

doc.add_heading('3.3 大三上学期（2026年9月-12月）', level=2)
semester_tasks = [
    '持续更新视频，累计达到10-20个',
    '完善个人网站，增加课件下载和推导说明',
    '整理项目日志，每周记录进展、问题、解决方案',
    '收集B站播放量、弹幕、评论数据',
    'Uniapp APP开发（如决定推进）：Vue基础→参数化交互组件→本地视频打包',
    '开始准备大创申报书初稿，用DeepSeek辅助生成',
    'GitHub贡献图积累到几百次提交',
]
for t in semester_tasks:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('3.4 大三下学期（2027年3月-4月）', level=2)
final_tasks = [
    '正式报名大创，此时已有10-20个视频和成熟网站',
    '答辩准备：笔记本展示GitHub代码+手机展示APP操作',
    '准备答辩PPT大纲（15页），包含背景、技术路线、成果、数据、后续计划',
    '完善申报书中的关键词：跨平台开发、程序化动画引擎、AI辅助编程、参数化可视化',
]
for f in final_tasks:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ============ 四、学习路径规划 ============
doc.add_heading('四、学习路径规划', level=1)

doc.add_heading('4.1 Python基础（第1-2周，每天3h）', level=2)
doc.add_paragraph('原则：不要从计算机专业教材学起，只学Manim用得到的部分。')
python_topics = [
    '变量、循环、判断、函数、类的基本概念',
    'numpy基础，会用数组和简单运算',
    '用VS Code或PyCharm搭建开发环境',
    '推荐资料：廖雪峰Python教程 + B站小甲鱼零基础Python',
]
for t in python_topics:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('4.2 Manim入门（第3-4周，每天3h）', level=2)
manim_topics = [
    '安装ManimCE（不要用旧的3b1b分支）',
    '学会场景（Scene）、Mobject动画流程',
    '从他人的开源代码改起：改颜色、速度、文字',
    '目标：产出一个30秒的运动学动画（如匀速直线运动的小车），生成mp4',
    '核心资源：ManimCE官方中文文档、B站「亲手画出动态数学」入门课、B站「图灵与欧拉的manim课程」',
    '特别注意：只用ManimCE版本，不要用旧版manimgl或3b1b/manim，否则代码会报错',
]
for t in manim_topics:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('4.3 Manim深入（第5-6周，每天3h）', level=2)
manim_advanced = [
    '函数图像：Axes + ParametricFunction',
    '动点/轨迹：ValueTracker + always_redraw',
    '物理：匀变速运动的位置-时间图、速度-时间图，运动分解',
    '配合DeepSeek：把场景需求喂给API，让它生成初稿，你来调试渲染',
    '建立代码模板库：坐标系模板、函数动画模板、文字说明模板',
]
for t in manim_advanced:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('4.4 DeepSeek API学习（第4周起融入）', level=2)
api_topics = [
    '安装Cherry Studio或VS Code Continue插件',
    '配置DeepSeek API Key（100元预算，足够支撑整个项目周期）',
    '学习标准Prompt模板的使用',
    '省钱技巧：日常简单问题用DeepSeek网页版免费功能，API只在生成长代码时使用',
    '110元约等于2.2亿Token，保守估计足够生成调试完20个视频的全部代码',
]
for t in api_topics:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('4.5 HTML/CSS速成（第5-6周穿插，共6h）', level=2)
html_topics = [
    'HTML骨架（1h）：html, head, body, div, a, img, iframe',
    'CSS核心（2h）：选择器（类、ID）、盒模型、display:flex布局、margin/padding',
    '网页排版实战（2h）：跟着MDN的"构建一个网页"教程做',
    'B站视频嵌入（0.5h）：学会iframe嵌入B站播放器',
    '响应式初步（0.5h）：了解@media查询',
]
for t in html_topics:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('4.6 Uniapp/APP开发（可选，暑假末+大三上）', level=2)
uniapp_topics = [
    'Uniapp基础（3天）：理解pages、组件、路由跳转，B站"黑马程序员Uniapp教程"前20集',
    'Vue基础（2天）：理解v-model、v-for、数据绑定',
    '用Claude Code生成APP代码：描述页面需求→Claude写Vue文件→HBuilder X预览→微调',
    '打包与发布（1天）：HBuilder X一键打包APK',
]
for t in uniapp_topics:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('4.7 核心学习资源', level=2)
resources_table = doc.add_table(rows=9, cols=2)
resources_table.style = 'Light Grid Accent 1'
resources_table.alignment = WD_TABLE_ALIGNMENT.CENTER
resources_table.rows[0].cells[0].text = '类别'
resources_table.rows[0].cells[1].text = '资源'

resources_data = [
    ['官方文档', 'ManimCE官方中文文档（必读）: docs.manim.cn'],
    ['在线演练', 'Manim交互式在线演练场: try.manim.community'],
    ['入门视频', 'B站「亲手画出动态数学」模块一至三'],
    ['进阶视频', 'B站「图灵与欧拉的manim课程」「Manim基础课」系列'],
    ['Python教程', '廖雪峰Python教程 + B站小甲鱼零基础Python'],
    ['社区问答', 'Manim官方Discord、Stack Overflow、Reddit r/manim'],
    ['中文社区', 'manim-kindergarten、乐正垂星B站、cai-hust学习笔记'],
    ['网页开发', 'MDN Web Docs "构建一个网页"教程'],
]
for i, row_data in enumerate(resources_data):
    for j, cell_data in enumerate(row_data):
        resources_table.rows[i+1].cells[j].text = cell_data

doc.add_page_break()

# ============ 五、团队分工 ============
doc.add_heading('五、团队分工与协作方案', level=1)

doc.add_heading('5.1 精准分工表', level=2)
team_table = doc.add_table(rows=6, cols=3)
team_table.style = 'Light Grid Accent 1'
team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
team_table.rows[0].cells[0].text = '成员'
team_table.rows[0].cells[1].text = '专业背景'
team_table.rows[0].cells[2].text = '核心职责'

team_data = [
    ['你（负责人）', '物理学（师范）', 'Manim动画制作、配音、AI辅助编程、项目统筹'],
    ['队员A', '数学（师范）', '知识点筛选、分镜脚本撰写、视频脚本校对'],
    ['队员B', '计算机科学', 'Python环境搭建、代码调试、批渲染、技术杂活'],
    ['队员C', '网络工程', 'GitHub Pages网站搭建、Uniapp环境配置、打包发布'],
    ['队员D', '暂定', '视频封面设计、B站标题标签优化、Canva使用'],
]
for i, row_data in enumerate(team_data):
    for j, cell_data in enumerate(row_data):
        team_table.rows[i+1].cells[j].text = cell_data

doc.add_heading('5.2 协作原则', level=2)
collab_rules = [
    '你的时间不要花在教他们上，而是给出明确的任务清单，让他们自学后来找你验收',
    '队员A立刻启动，暑假前就能拿到10个脚本',
    '队员C暑假自学GitHub Pages部署流程和Hugo模板使用',
    '队员B暑假自学Python基础+Manim安装，达到能帮你调试代码、批渲染的水平',
    'D负责所有视频封面制作、B站标题和标签优化',
    '每周日晚开30分钟线上站会，核对进度并调整下周任务',
]
for r in collab_rules:
    doc.add_paragraph(r, style='List Bullet')

doc.add_page_break()

# ============ 六、注意事项 ============
doc.add_heading('六、关键注意事项与避坑指南', level=1)

doc.add_heading('6.1 环境与工具', level=2)
env_notes = [
    '只认准Manim Community Edition (ManimCE)，不要用旧的3b1b分支或manimgl版本',
    '渲染对显卡不是必须的，但有独立显卡会极大提升效率',
    '优先用Cairo渲染器（默认），有独显可用OpenGL加速（命令：manim -pql --renderer=opengl）',
    '先用笔记本跑一遍安装流程，渲染简单动画测试，速度能接受就接着用',
    '用 -ql 参数出低画质预览看效果，最后用 -qh 出成品',
    '如果电脑带不动，可考虑短期租用AutoDL等云GPU服务',
    'Mac用户建议直接用Cairo版本',
    '显卡驱动一定要更新',
]
for n in env_notes:
    doc.add_paragraph(n, style='List Bullet')

doc.add_heading('6.2 学习过程', level=2)
learn_notes = [
    '不要从计算机专业教材学起Python，只学Manim用得到的部分',
    'Manim从他人的开源代码改起，先改参数改文字，再慢慢写原创',
    '遇到卡点超过2小时还没解决，立刻上官方Discord或GitHub Discussions去问',
    '每天3-4小时高效学习远胜于8小时低效，劳逸结合是持久战',
    '用DeepSeek网页版（免费）先行检验，积累好prompt经验再用API',
    '录制工作过程：偶尔录屏记录调试到渲染成功的过程，可剪辑成"幕后花絮"',
    '先用现有知识换新技术，不要被APP开发的难度吓到',
]
for n in learn_notes:
    doc.add_paragraph(n, style='List Bullet')

doc.add_heading('6.3 项目开发', level=2)
dev_notes = [
    '代码备份：每天用Git提交，哪怕只写一行注释也要推上去',
    '每个视频严格控制在1-3分钟，只讲1个概念',
    '脚本先行：先写画外音文案，再写代码，最后录音合成',
    '统一配色方案、字体、动画节奏，写入全局配置文件',
    '建立公共素材库：常用符号SVG、背景音乐片段、片头片尾动画',
    '遇到报错不要硬扛，把错误信息喂给AI让它修改',
    '先让滑块能滑、画面能动就算成功，降低初始审美标准',
]
for n in dev_notes:
    doc.add_paragraph(n, style='List Bullet')

doc.add_heading('6.4 大创申报', level=2)
apply_notes = [
    '定位成"技术开发类"项目，不要过多强调"做了多少课件"',
    '关键词：跨平台开发、程序化动画引擎、AI辅助编程、参数化可视化',
    '多提"响应国家教育数字化转型战略"、"探索AI时代师范生数字素养提升路径"',
    '找一位对教育技术、AI或新媒体感兴趣的青年教师（副教授以下）',
    '带着简易Demo去谈指导老师，效果远胜空手描述',
    'GitHub提交记录是任何评委都无法质疑的"过程证据"',
    '答辩时笔记本展示GitHub代码+手机展示APP操作，冲击力远超PPT',
    '在暑假末就开始整理项目日志（Word文档），这就是申报书和中期报告的素材',
]
for n in apply_notes:
    doc.add_paragraph(n, style='List Bullet')

doc.add_heading('6.5 B站运营', level=2)
bilibili_notes = [
    '合集名：如"看得懂的物理/数学｜Manim动画系列"',
    '每集标题：知识点+一句话吸引力，如"3分钟看懂：二次函数的a,b,c到底在控制什么"',
    '播放量互动：让身边同学、课程老师集中观看和评论，快速积累最初1000播放',
    '简介区放置网站链接，写"更多推导过程和课件下载见评论区/网站"',
    '大创答辩时直接展示B站后台的播放量、粉丝数、弹幕热词',
]
for n in bilibili_notes:
    doc.add_paragraph(n, style='List Bullet')

doc.add_heading('6.6 DeepSeek API使用', level=2)
api_notes = [
    '100元预算+官方送10元体验金=110元，足够整个项目周期',
    '日常简单问题用DeepSeek网页版（免费），不要浪费API额度',
    'API只在需要生成长代码、复杂脚本、批量处理时使用',
    '生成代码时先让它写核心逻辑，你手动加注释和微调，减少反复修改次数',
    '一个API Key可以在多台电脑上使用，团队共用一个预算',
    '不要把API Key泄露到公开网页上',
]
for n in api_notes:
    doc.add_paragraph(n, style='List Bullet')

doc.add_page_break()

# ============ 七、DeepSeek API 使用策略 ============
doc.add_heading('七、DeepSeek API 使用策略', level=1)

doc.add_heading('7.1 四大核心用途', level=2)
api_uses = [
    'Manim代码生成与调试（核心用途）：自然语言转代码、代码解释与教学、批量生成模板',
    '个人网站HTML/CSS/JS代码生成：页面结构、调试修改、功能代码片段',
    '物理/数学脚本撰写辅助：知识点拆解、公式转画外音',
    '大创相关文本辅助：申报书草拟、答辩PPT大纲',
]
for u in api_uses:
    doc.add_paragraph(u, style='List Number')

doc.add_heading('7.2 标准Prompt模板（保存反复使用）', level=2)
prompt_template = (
    "你是一个精通ManimCE的数学动画专家。请根据以下描述生成可直接运行的Python代码。\n\n"
    "描述：[详细描述场景，包括坐标系、对象、运动、颜色、标注、相机移动等]\n\n"
    "要求：\n"
    "1. 使用ManimCE 0.18+语法，所有import齐全。\n"
    "2. 关键步骤添加中文注释。\n"
    "3. 渲染命令默认为 manim -pql file.py SceneName。\n"
    "如果后续报错，我将粘贴错误信息，请修正后输出完整代码。"
)
p = doc.add_paragraph()
run = p.add_run(prompt_template)
run.font.size = Pt(10)
run.font.name = 'Consolas'

doc.add_heading('7.3 费用估算', level=2)
cost_items = [
    '1次完整的"生成动画代码+1-2轮修改"：约2-5万Token',
    '110元约等于2.2亿Token',
    '保守估计：足够支撑生成调试完20个视频的全部代码+网页代码+申报书文本',
    '省钱技巧：日常问题用免费网页版，API仅用于复杂代码生成',
]
for c in cost_items:
    doc.add_paragraph(c, style='List Bullet')

doc.add_page_break()

# ============ 八、大创申报策略 ============
doc.add_heading('八、大创申报策略与技巧', level=1)

doc.add_heading('8.1 项目名称与类别', level=2)
doc.add_paragraph('项目名称：基于Manim与AI开发的中学生科学可视化互动学习APP')
doc.add_paragraph('项目类别：创新训练项目')

doc.add_heading('8.2 五大过审策略', level=2)
strategies = [
    '策略1：用"技术创新"标签替代"教学实践"标签 → 关键词用跨平台开发、程序化动画引擎、AI辅助编程',
    '策略2：用GitHub开源和真实用户数据增加说服力 → 几百次提交+5000播放+30条含金量评论',
    '策略3：规避"敏感词"，拥抱"政策红利" → 多提教育数字化、AI赋能、自主探究、开源共享',
    '策略4：提前搞定指导老师 → 找对教育技术感兴趣的青年教师，带着Demo去谈',
    '策略5：答辩时直接掏出"手机+笔记本" → 笔记本展示代码仓库，手机打开APP让评委操作',
]
for s in strategies:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('8.3 申报书核心内容框架', level=2)
framework = [
    '项目背景：中学抽象概念理解痛点 + 数字化教学资源建设趋势',
    '技术路线：内容选题→脚本设计→AI辅助动画制作→视频合成→双载体发布→反馈迭代',
    '现有成果（申报时）：10-20个视频 + B站数据 + 网站 + GitHub仓库',
    '创新点：技术融合+双载体传播+交互模式创新+学科交叉',
    '预期成果：视频+网站+开源代码+论文',
    '推广数据：B站播放量、弹幕、评论',
    '后续计划：APP迭代、更多视频产出、开源社区建设',
]
for f in framework:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ============ 九、每日作息 ============
doc.add_heading('九、每日作息与学习节奏建议', level=1)

doc.add_heading('9.1 暑假每日安排', level=2)
daily_items = [
    '上午（2h）：学习新知识（Python/Manim/HTML），看教程记笔记',
    '下午（2h）：实践项目，写代码、渲染动画、与AI对话迭代',
    '晚上（0.5h）：复盘今天遇到的问题，记录到GitHub Issues或笔记本；push代码',
    '周末：选一天完全休息，另一天可以自由创造或处理团队协作任务',
]
for d in daily_items:
    doc.add_paragraph(d, style='List Bullet')

doc.add_heading('9.2 学习节奏原则', level=2)
rhythm = [
    '每天3-4小时高效学习远胜于8小时低效',
    '运动、社交保持正常，这是持久战',
    '遇到卡点不要硬扛，超过2小时就去社区问',
    '用AI加速学习：看不懂的代码问AI，报错贴给AI，让AI出练习题',
    '先跑通流程再追求完美，降低初始标准',
]
for r in rhythm:
    doc.add_paragraph(r, style='List Bullet')

doc.add_page_break()

# ============ 十、阶段性检查清单 ============
doc.add_heading('十、阶段性检查清单', level=1)

doc.add_heading('10.1 暑假前（环境准备）', level=2)
pre_checks = [
    '☐ B帮你装好Python环境和ManimCE，最简单的示例渲染成功',
    '☐ 注册B站账号和GitHub Organization，把团队加到GitHub仓库',
    '☐ C先搭一个空白GitHub Pages，放一句"视频即将上线"',
    '☐ 你自己跑通"Python打印Hello World→numpy画图→Manim画一个正方形"',
    '☐ 用Cherry Studio配置好DeepSeek API Key',
    '☐ 确定前10个知识点，你与A一起从人教版必修一里挑',
]
for c in pre_checks:
    doc.add_paragraph(c)

doc.add_heading('10.2 暑假第2周末', level=2)
week2_checks = [
    '☐ 能独立写一个生成物理运动轨迹数据的Python脚本',
    '☐ 能在AI帮助下解读Manim官方例子的代码',
    '☐ Python基础概念（变量/循环/函数/类）基本掌握',
]
for c in week2_checks:
    doc.add_paragraph(c)

doc.add_heading('10.3 暑假第4周末', level=2)
week4_checks = [
    '☐ 有一个渲染成功的.mp4文件，包含坐标系、运动的点、矢量箭头',
    '☐ 能用标准Prompt模板让DeepSeek生成Manim代码',
    '☐ Manim基本元素和动画模式已掌握',
]
for c in week4_checks:
    doc.add_paragraph(c)

doc.add_heading('10.4 暑假结束', level=2)
end_checks = [
    '☐ 6-8个已上传B站的教学视频',
    '☐ 个人网站可以访问，每个视频有独立介绍页面和嵌入播放器',
    '☐ GitHub仓库有完整的代码和课件，提交记录丰富',
    '☐ 队员A已提供至少10个视频的分镜脚本',
    '☐ HTML/CSS已入门，能看懂标签、class、id概念',
    '☐ 大创申报书草稿已开始撰写',
]
for c in end_checks:
    doc.add_paragraph(c)

doc.add_heading('10.5 大三上学期结束', level=2)
semester_checks = [
    '☐ 累计10-20个视频，B站合集总播放量目标≥2万',
    '☐ 个人网站完善，课件下载功能可用',
    '☐ GitHub提交记录几百次',
    '☐ 项目日志完整（每周进展、问题、解决方案）',
    '☐ APP原型Demo可运行（如决定推进）',
    '☐ 大创申报书初稿完成',
]
for c in semester_checks:
    doc.add_paragraph(c)

doc.add_page_break()

# ============ 结语 ============
doc.add_heading('结语', level=1)
doc.add_paragraph(
    '这个项目如果扎实推进到暑假结束，大三下申报大创时你几乎是"带着成品答辩"。'
    '上手门槛在Manim，但你是师范物理，这恰恰是你将来的教学硬技能，非常值得投入。'
)
doc.add_paragraph(
    '原路线（纯视频+网站）= 可能被打下来的"教育实践类"项目；'
    '升级路线（互动APP+视频引流+网站沉淀）= 容易通过的"教育技术创新类"项目。'
    '核心技术不变（Manim+AI辅助），但载体升级，评审感知完全不同。'
)
doc.add_paragraph(
    '你现在可以立刻开始的三个动作：\n'
    '1. 用PyCharm跑通Manim的SquareToCircle，确保环境无误。\n'
    '2. 在Claude Code里用标准Prompt生成第一个"匀加速运动"的代码，看看效果。\n'
    '3. 和队员A现在就用在线文档，开始罗列20个知识点的清单。'
)
p_final = doc.add_paragraph()
run = p_final.add_run('这个暑假，你不仅能做出让自己惊喜的作品，更会建立一套终身受用的"AI+X"创作能力。祝你暑假一战成名！')
run.bold = True

# 保存
output_path = r'D:\pythonapp\open\大创项目规划_完整版.docx'
doc.save(output_path)
print(f'文档已保存到: {output_path}')
