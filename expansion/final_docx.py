#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""最终版：封面 + 可跳转目录 + 页码 + 页眉页脚"""

import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    doc.add_page_break()

def create_cover(doc):
    """创建封面"""
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('灵魂的空壳')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('现存"社会主义"国家与经典定义的彻底断裂')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run('━' * 30)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run('一项结合马列毛经典与西方马克思主义的综合批判')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for _ in range(4):
        doc.add_paragraph()

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run('基于马克思、恩格斯、列宁、毛泽东经典著作')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    author2 = doc.add_paragraph()
    author2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author2.add_run('及葛兰西、马尔库塞、卢卡奇、阿尔都塞、普兰查斯等西方马克思主义思想家的批判理论')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    for _ in range(4):
        doc.add_paragraph()

    year = doc.add_paragraph()
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = year.add_run('二〇二五年')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    add_page_break(doc)

def create_toc(doc):
    """创建可跳转的目录"""
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run('目  录')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    doc.add_paragraph()

    # 插入TOC域代码 - 这会在Word中生成可跳转的目录
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar2)

    # 目录内容（占位文本）
    toc_items = [
        '引言：社会主义的星辰与泥沼',
        '第一篇 理论基石',
        '  第一章 经典社会主义的完整定义：被背叛的承诺',
        '第二篇 全球社会主义运动的历史轨迹',
        '  第二章 社会主义运动的全球历史轨迹',
        '  第三章 苏联的兴亡：从十月革命到红旗落地',
        '第三篇 四面镜子：四国分析',
        '  第四章 朝鲜：封建世袭的"朝显王国"',
        '  第五章 越南：官僚资本主义的"戏仿"',
        '  第六章 老挝：被遗忘的"社会主义原始积累"',
        '  第七章 古巴：冻结在冷战中的"革命化石"',
        '第四篇 最复杂的理论案例：中国',
        '  第八章 中国：社会主义的"变体"还是"背叛"？',
        '第五篇 六大维度的系统比较',
        '  第九章 权力结构的比较分析',
        '  第十章 经济基础的比较分析',
        '  第十一章 意识形态的比较分析',
        '  第十二章 社会结构与阶级的比较分析',
        '  第十三章 法治与公民权利的比较分析',
        '  第十四章 文化道德与社会风气的比较分析',
        '第六篇 新议题',
        '  第十五章 "社会主义"标签的解构',
        '  第十六章 帝国主义、全球化与依附关系',
        '  第十七章 数字时代的新形态控制',
        '第七篇 综合判断与理论反思',
        '  第十八章 灵魂的空壳',
        '  第十九章 毛泽东晚年思想的当代启示',
        '  第二十章 21世纪社会主义的可能出路',
        '第二十一章 数字时代的社会主义',
        '第二十二章 生态社会主义',
        '第二十三章 性别与社会主义',
        '第二十四章 少数民族与社会主义',
        '第二十五章 21世纪社会主义的理论重构',
        '第二十六章 全球左翼运动的当代图景',
        '第二十七章 社会主义与人权',
        '第二十八章 社会主义的未来',
        '第二十九章 马克思主义中国化的理论演进',
        '第三十章 社会主义与市场经济',
        '第三十一章 社会主义与全球化',
        '第三十二章 社会主义与科技革命',
        '第三十三章 灵魂的重生：典型社会主义国家的构想',
        '第三十四章 新左派的自我修养',
        '结语：没有答案的审判与尚存希望的未来',
        '参考文献',
    ]

    for item in toc_items:
        run4 = paragraph.add_run()
        run4.text = item + '\n'
        run4.font.size = Pt(11)

    run5 = paragraph.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar3)

    # 添加提示文本
    hint = doc.add_paragraph()
    hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hint.add_run('（在Word中右键点击目录选择"更新域"以生成完整目录）')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True

    add_page_break(doc)

def add_page_numbers(doc):
    """添加页码到页脚"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加"— "前缀
        run1 = fp.add_run('— ')
        run1.font.size = Pt(9)
        run1.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # 插入页码域
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run1._r.addnext(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        fldChar1.addnext(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        instrText.addnext(fldChar2)

        # 添加"— "后缀
        run2 = fp.add_run(' —')
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

def add_headers(doc):
    """添加页眉"""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = hp.add_run('灵魂的空壳：现存"社会主义"国家与经典定义的彻底断裂')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.italic = True

        # 页眉下划线
        pPr = hp._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '999999')
        pBdr.append(bottom)
        pPr.append(pBdr)

def main():
    # 读取现有docx
    input_path = r'D:\pythonapp\open\expansion\灵魂的空壳_扩写版.docx'
    output_path = r'D:\pythonapp\open\expansion\灵魂的空壳_完整版.docx'

    print("读取现有文档...")
    existing_doc = Document(input_path)

    # 创建新文档
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # 设置页面
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    # 创建封面
    print("创建封面...")
    create_cover(doc)

    # 创建目录
    print("创建目录...")
    create_toc(doc)

    # 复制正文
    print("复制正文内容...")
    for para in existing_doc.paragraphs:
        new_para = doc.add_paragraph()
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.bold:
                new_run.font.bold = run.font.bold
            if run.font.italic:
                new_run.font.italic = run.font.italic
            if run.font.color and run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
        if para.alignment:
            new_para.alignment = para.alignment
        if para.style.name.startswith('Heading'):
            try:
                new_para.style = doc.styles[para.style.name]
            except:
                pass

    # 添加页眉
    print("添加页眉...")
    add_headers(doc)

    # 添加页码
    print("添加页码...")
    add_page_numbers(doc)

    # 保存
    print(f"保存至: {output_path}")
    doc.save(output_path)
    print("完成！")

if __name__ == '__main__':
    main()
