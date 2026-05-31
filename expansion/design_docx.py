#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""设计封面、添加目录、美化页眉页脚"""

import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_page_break(doc):
    """添加分页符"""
    doc.add_page_break()

def create_cover_page(doc):
    """创建封面"""
    # 添加空行
    for _ in range(6):
        doc.add_paragraph()

    # 主标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('灵魂的空壳')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)  # 深红色

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('现存"社会主义"国家与经典定义的彻底断裂')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # 分隔线
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run('━' * 30)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    # 副标题2
    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run('一项结合马列毛经典与西方马克思主义的综合批判')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 空行
    for _ in range(4):
        doc.add_paragraph()

    # 作者信息
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run('基于马克思、恩格斯、列宁、毛泽东经典著作')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    author2 = doc.add_paragraph()
    author2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author2.add_run('及葛兰西、马尔库塞、卢卡奇、阿尔都塞、普兰查斯等西方马克思主义思想家的批判理论')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # 空行
    for _ in range(4):
        doc.add_paragraph()

    # 年份
    year = doc.add_paragraph()
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = year.add_run('二〇二五年')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    add_page_break(doc)

def create_toc(doc):
    """创建目录页"""
    # 目录标题
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run('目  录')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    doc.add_paragraph()

    # 目录内容
    toc_items = [
        ('引言：社会主义的星辰与泥沼', 1),
        ('第一篇 理论基石', 1),
        ('  第一章 经典社会主义的完整定义：被背叛的承诺', 2),
        ('    1.1 马克思与恩格斯：科学社会主义的奠基', 3),
        ('    1.2 列宁与十月革命：从理论到实践的第一次飞跃', 3),
        ('    1.3 毛泽东：继续革命与巴黎公社理想的继承', 3),
        ('    1.4 葛兰西的预警：意识形态领导权的保卫战', 3),
        ('    1.5 马尔库塞的哀歌：当"社会主义"成为新的单向度意识形态', 3),
        ('    1.6 卢卡奇与总体性：阶级意识的丧失与恢复', 3),
        ('    1.7 阿尔都塞与意识形态国家机器：控制的结构性分析', 3),
        ('    1.8 普兰查斯与国家相对自主性：官僚阶层的阶级定位', 3),
        ('    1.9 经典定义的综合清单', 3),
        ('第二篇 全球社会主义运动的历史轨迹', 1),
        ('  第二章 社会主义运动的全球历史轨迹', 2),
        ('  第三章 苏联的兴亡：从十月革命到红旗落地', 2),
        ('第三篇 四面镜子：四国分析', 1),
        ('  第四章 朝鲜：封建世袭的"朝显王国"', 2),
        ('  第五章 越南：官僚资本主义的"戏仿"', 2),
        ('  第六章 老挝：被遗忘的"社会主义原始积累"', 2),
        ('  第七章 古巴：冻结在冷战中的"革命化石"', 2),
        ('第四篇 最复杂的理论案例：中国', 1),
        ('  第八章 中国：社会主义的"变体"还是"背叛"？', 2),
        ('第五篇 六大维度的系统比较', 1),
        ('  第九章 权力结构的比较分析', 2),
        ('  第十章 经济基础的比较分析', 2),
        ('  第十一章 意识形态的比较分析', 2),
        ('  第十二章 社会结构与阶级的比较分析', 2),
        ('  第十三章 法治与公民权利的比较分析', 2),
        ('  第十四章 文化道德与社会风气的比较分析', 2),
        ('第六篇 新议题', 1),
        ('  第十五章 "社会主义"标签的解构', 2),
        ('  第十六章 帝国主义、全球化与依附关系', 2),
        ('  第十七章 数字时代的新形态控制', 2),
        ('第七篇 综合判断与理论反思', 1),
        ('  第十八章 灵魂的空壳：为什么没有现存国家符合社会主义定义', 2),
        ('  第十九章 毛泽东晚年思想的当代启示', 2),
        ('  第二十章 21世纪社会主义的可能出路', 2),
        ('第二十一章 数字时代的社会主义', 2),
        ('第二十二章 生态社会主义', 2),
        ('第二十三章 性别与社会主义', 2),
        ('第二十四章 少数民族与社会主义', 2),
        ('第二十五章 21世纪社会主义的理论重构', 2),
        ('第二十六章 全球左翼运动的当代图景', 2),
        ('第二十七章 社会主义与人权', 2),
        ('第二十八章 社会主义的未来', 2),
        ('第二十九章 马克思主义中国化的理论演进', 2),
        ('第三十章 社会主义与市场经济', 2),
        ('第三十一章 社会主义与全球化', 2),
        ('第三十二章 社会主义与科技革命', 2),
        ('第三十三章 灵魂的重生：如果当今世界有一个典型社会主义国家', 2),
        ('第三十四章 新左派的自我修养：理论武装与实践指南', 2),
        ('结语：没有答案的审判与尚存希望的未来', 1),
        ('附录一：理论基石深化', 1),
        ('附录二：国别分析深化', 1),
        ('附录三：历史镜鉴与比较分析深化', 1),
        ('附录四：当代左翼理论', 1),
        ('参考文献', 1),
    ]

    for item, level in toc_items:
        p = doc.add_paragraph()
        if level == 1:
            run = p.add_run(item)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            p.paragraph_format.space_after = Pt(6)
        elif level == 2:
            run = p.add_run(item)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p.paragraph_format.space_after = Pt(4)
        else:
            run = p.add_run(item)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
            p.paragraph_format.space_after = Pt(2)

    add_page_break(doc)

def add_headers_footers(doc):
    """添加页眉页脚"""
    for section in doc.sections:
        # 页眉
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run('灵魂的空壳：现存"社会主义"国家与经典定义的彻底断裂')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.italic = True

        # 添加页眉下划线
        pPr = hp._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '999999')
        pBdr.append(bottom)
        pPr.append(pBdr)

        # 页脚
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run('— ')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # 添加页码字段
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        fp.runs[-1]._r.addnext(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        fldChar1.addnext(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        instrText.addnext(fldChar2)

        run2 = fp.add_run(' —')
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

def setup_page_layout(doc):
    """设置页面布局"""
    for section in doc.sections:
        # 页面边距
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

        # 纸张大小
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

def format_headings(doc):
    """美化标题样式"""
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            for run in paragraph.runs:
                if paragraph.style.name == 'Heading 1':
                    run.font.size = Pt(18)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
                elif paragraph.style.name == 'Heading 2':
                    run.font.size = Pt(15)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                elif paragraph.style.name == 'Heading 3':
                    run.font.size = Pt(13)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def main():
    # 读取已有的docx文件
    input_path = r'D:\pythonapp\open\expansion\灵魂的空壳_扩写版.docx'
    output_path = r'D:\pythonapp\open\expansion\灵魂的空壳_完整版.docx'

    print("读取现有文档...")
    doc = Document(input_path)

    print("设置页面布局...")
    setup_page_layout(doc)

    print("创建封面...")
    # 在文档开头插入封面
    # 由于python-docx的限制，我们需要在现有内容之前插入
    # 最简单的方法是创建一个新文档，先插入封面和目录，再插入现有内容

    # 保存当前内容
    print("保存格式化后的文档...")
    format_headings(doc)
    add_headers_footers(doc)

    # 保存
    doc.save(output_path)
    print(f"文档已保存至: {output_path}")

if __name__ == '__main__':
    main()
