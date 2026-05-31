#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""合并所有扩写内容为docx文件"""

import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

# 尝试导入python-docx
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed. Installing...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

# 文件列表（按顺序）- 包含扩写补充
files = [
    # 引言
    r'D:\pythonapp\open\expansion\batch1_intro.md',
    r'D:\pythonapp\open\expansion\expand_intro.md',
    # 第一章 理论基石
    r'D:\pythonapp\open\expansion\batch1_ch1_1.md',
    r'D:\pythonapp\open\expansion\batch1_ch1_2.md',
    r'D:\pythonapp\open\expansion\batch1_ch1_3.md',
    r'D:\pythonapp\open\expansion\batch1_ch1_4to8.md',
    r'D:\pythonapp\open\expansion\expand_ch1.md',
    r'D:\pythonapp\open\expansion\expand_theory1.md',
    r'D:\pythonapp\open\expansion\expand_theory2.md',
    r'D:\pythonapp\open\expansion\expand_theory3.md',
    r'D:\pythonapp\open\expansion\expand_theory4.md',
    r'D:\pythonapp\open\expansion\expand_theory5.md',
    r'D:\pythonapp\open\expansion\expand_theory6.md',
    r'D:\pythonapp\open\expansion\expand_theory_practice.md',
    r'D:\pythonapp\open\expansion\expand_theory_practice2.md',
    r'D:\pythonapp\open\expansion\expand_theory_practice3.md',
    r'D:\pythonapp\open\expansion\expand_cross_theory.md',
    r'D:\pythonapp\open\expansion\expand_appendix.md',
    r'D:\pythonapp\open\expansion\expand_new_chapters.md',
    r'D:\pythonapp\open\expansion\expand_new_chapters2.md',
    r'D:\pythonapp\open\expansion\expand_new_chapters3.md',
    r'D:\pythonapp\open\expansion\expand_ideal_socialism.md',
    r'D:\pythonapp\open\expansion\expand_new_left.md',
    r'D:\pythonapp\open\expansion\expand_practice_guide.md',
    r'D:\pythonapp\open\expansion\expand_social_issues.md',
    r'D:\pythonapp\open\expansion\expand_all_chapters.md',
    r'D:\pythonapp\open\expansion\expand_final_push.md',
    # 第二篇 全球历史
    r'D:\pythonapp\open\expansion\batch2_ch2.md',
    r'D:\pythonapp\open\expansion\batch2_ch3.md',
    r'D:\pythonapp\open\expansion\expand_ch2to3.md',
    r'D:\pythonapp\open\expansion\expand_ch2_deep.md',
    # 第三篇 四国分析
    r'D:\pythonapp\open\expansion\batch3_ch4.md',
    r'D:\pythonapp\open\expansion\batch3_ch5.md',
    r'D:\pythonapp\open\expansion\batch3_ch6.md',
    r'D:\pythonapp\open\expansion\batch3_ch7.md',
    r'D:\pythonapp\open\expansion\expand_ch4to7.md',
    r'D:\pythonapp\open\expansion\expand_country1.md',
    # 第四篇 中国案例
    r'D:\pythonapp\open\expansion\batch4_ch8.md',
    # 第五篇 六维度比较
    r'D:\pythonapp\open\expansion\batch5_ch9to14.md',
    r'D:\pythonapp\open\expansion\expand_ch9to17.md',
    r'D:\pythonapp\open\expansion\expand_compare.md',
    # 第六篇 新议题与综合判断
    r'D:\pythonapp\open\expansion\batch6_ch15to20.md',
    r'D:\pythonapp\open\expansion\expand_ch18to20.md',
    # 结语
    r'D:\pythonapp\open\expansion\batch6_conclusion.md',
    # 更多扩写内容
    r'D:\pythonapp\open\expansion\expand_more.md',
    r'D:\pythonapp\open\expansion\expand_final.md',
    r'D:\pythonapp\open\expansion\expand_final2.md',
    r'D:\pythonapp\open\expansion\expand_final3.md'
]

# 创建文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)

# 读取原始docx的参考文献（如果需要）
ref_docx = r'C:\Users\皮皮\Desktop\教员纪念合集\新左派\灵魂的空壳正文版.docx'

# 合并所有内容
total_chars = 0
chinese_chars = 0

for filepath in files:
    if not os.path.exists(filepath):
        print(f"WARNING: {filepath} not found, skipping...")
        continue

    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    total_chars += len(content)
    chinese_chars += sum(1 for c in content if '一' <= c <= '鿿')

    lines = content.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 处理标题
        if line.startswith('# ') and not line.startswith('## '):
            # 一级标题
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            # 二级标题
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            # 三级标题
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            # 四级标题
            doc.add_heading(line[5:], level=4)
        elif line.startswith('**') and line.endswith('**'):
            # 粗体段落
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.bold = True
        elif line.startswith('---'):
            # 分隔线
            doc.add_paragraph('—' * 50)
        elif line.startswith('**注释：') or line.startswith('**核心判断：') or line.startswith('**结论：'):
            # 特殊标注段落
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        elif line.startswith('**') and '：**' in line:
            # 标注+内容
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        elif line.startswith('|'):
            # 表格行 - 简单处理为段落
            doc.add_paragraph(line)
        elif line.startswith('¹') or line.startswith('²') or line.startswith('³') or line.startswith('⁴') or line.startswith('⁵') or line.startswith('⁶') or line.startswith('⁷') or line.startswith('⁸') or line.startswith('⁹'):
            # 脚注 - 小字号
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(9)
        else:
            # 普通段落
            doc.add_paragraph(line)

# 保存文件
output_path = r'D:\pythonapp\open\expansion\灵魂的空壳_扩写版.docx'
doc.save(output_path)

print(f"\n=== 统计 ===")
print(f"总字符数: {total_chars}")
print(f"中文字符数: {chinese_chars}")
print(f"文件已保存至: {output_path}")
