# -*- coding: utf-8 -*-
"""Convert Unicode sub/sup/frac chars to OMath in any document."""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

SRC = r'c:\Users\皮皮\Downloads\物理实验系统教程_带安装图.docx'
DST = r'c:\Users\皮皮\Downloads\物理实验系统教程_带安装图_OMath版.docx'

OMATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NSMAP = {'m': OMATH_NS, 'w': W_NS}

# Unicode subscript digits: ₀₁₂₃₄₅₆₇₈₉ (U+2080-U+2089)
_SUB_DIGITS = '₀₁₂₃₄₅₆₇₈₉'
SUB_MAP = {}
for i, ch in enumerate(_SUB_DIGITS):
    SUB_MAP[ch] = str(i)
SUB_MAP['ᵢ'] = 'i'  # ᵢ Latin subscript small letter i
SUB_MAP['ₙ'] = 'n'  # ₙ Latin subscript small letter n

# Unicode superscript digits: ⁰¹²³⁴⁵⁶⁷⁸⁹ (U+2070-U+2079)
_SUP_DIGITS = '⁰¹²³⁴⁵⁶⁷⁸⁹'
SUP_MAP = {}
for i, ch in enumerate(_SUP_DIGITS):
    SUP_MAP[ch] = str(i)
SUP_MAP['⁻'] = '-'  # ⁻ superscript minus

FRAC_MAP = {
    '½': ('1', '2'),  # 1/2
    '⅓': ('1', '3'),  # 1/3
    '⅔': ('2', '3'),  # 2/3
    '¼': ('1', '4'),  # 1/4
    '¾': ('3', '4'),  # 3/4
}

ALL_SPECIAL = set(SUB_MAP) | set(SUP_MAP) | set(FRAC_MAP)


def has_special(text):
    return any(ch in ALL_SPECIAL for ch in text)


def copy_rPr(src_run_elem):
    rPr = src_run_elem.find(qn('w:rPr'))
    if rPr is not None:
        return etree.fromstring(etree.tostring(rPr))
    return None


def make_text_run(text, rPr_copy=None):
    r = OxmlElement('w:r')
    if rPr_copy is not None:
        r.append(rPr_copy)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    return r


def _omath_r(text):
    r = etree.SubElement(etree.Element('dummy', nsmap=NSMAP), f'{{{OMATH_NS}}}r')
    # Add font declaration so Word knows to use Cambria Math
    rPr = etree.SubElement(r, f'{{{W_NS}}}rPr')
    rFonts = etree.SubElement(rPr, f'{{{W_NS}}}rFonts')
    rFonts.set(f'{{{W_NS}}}ascii', 'Cambria Math')
    rFonts.set(f'{{{W_NS}}}hAnsi', 'Cambria Math')
    t = etree.SubElement(r, f'{{{OMATH_NS}}}t')
    t.text = text
    t.set(f'{{{W_NS}}}space', 'preserve')
    return r


def make_omath_sub(base, sub_text):
    omath = etree.Element(f'{{{OMATH_NS}}}oMath', nsmap=NSMAP)
    sSub = etree.SubElement(omath, f'{{{OMATH_NS}}}sSub')
    e = etree.SubElement(sSub, f'{{{OMATH_NS}}}e')
    e.append(_omath_r(base))
    sub = etree.SubElement(sSub, f'{{{OMATH_NS}}}sub')
    sub.append(_omath_r(sub_text))
    return omath


def make_omath_sup(base, sup_text):
    omath = etree.Element(f'{{{OMATH_NS}}}oMath', nsmap=NSMAP)
    sSup = etree.SubElement(omath, f'{{{OMATH_NS}}}sSup')
    e = etree.SubElement(sSup, f'{{{OMATH_NS}}}e')
    e.append(_omath_r(base))
    sup = etree.SubElement(sSup, f'{{{OMATH_NS}}}sup')
    sup.append(_omath_r(sup_text))
    return omath


def make_omath_frac(num, den):
    omath = etree.Element(f'{{{OMATH_NS}}}oMath', nsmap=NSMAP)
    f = etree.SubElement(omath, f'{{{OMATH_NS}}}f')
    num_e = etree.SubElement(f, f'{{{OMATH_NS}}}num')
    num_e.append(_omath_r(num))
    den_e = etree.SubElement(f, f'{{{OMATH_NS}}}den')
    den_e.append(_omath_r(den))
    return omath


def parse_segments(text):
    segments = []
    buf = []
    i, n = 0, len(text)
    while i < n:
        ch = text[i]
        if ch in SUB_MAP:
            if buf:
                segments.append(('text', ''.join(buf)))
                buf = []
            sub_chars = [SUB_MAP[ch]]
            i += 1
            while i < n and text[i] in SUB_MAP:
                sub_chars.append(SUB_MAP[text[i]])
                i += 1
            if segments and segments[-1][0] == 'text' and segments[-1][1] and segments[-1][1][-1].isalpha():
                base = segments[-1][1][-1]
                segments[-1] = ('text', segments[-1][1][:-1])
                if not segments[-1][1]:
                    segments.pop()
                segments.append(('sub', base, ''.join(sub_chars)))
            else:
                segments.append(('sub_raw', ''.join(sub_chars)))
            continue
        elif ch in SUP_MAP:
            if buf:
                segments.append(('text', ''.join(buf)))
                buf = []
            sup_chars = [SUP_MAP[ch]]
            i += 1
            while i < n and text[i] in SUP_MAP:
                sup_chars.append(SUP_MAP[text[i]])
                i += 1
            if segments and segments[-1][0] == 'text' and segments[-1][1] and segments[-1][1][-1].isalpha():
                base = segments[-1][1][-1]
                segments[-1] = ('text', segments[-1][1][:-1])
                if not segments[-1][1]:
                    segments.pop()
                segments.append(('sup', base, ''.join(sup_chars)))
            else:
                segments.append(('sup_raw', ''.join(sup_chars)))
            continue
        elif ch in FRAC_MAP:
            if buf:
                segments.append(('text', ''.join(buf)))
                buf = []
            num, den = FRAC_MAP[ch]
            segments.append(('frac', num, den))
            i += 1
            continue
        else:
            buf.append(ch)
            i += 1
    if buf:
        segments.append(('text', ''.join(buf)))
    return segments


def merge_split_subscripts(para_elem):
    runs = list(para_elem.findall(qn('w:r')))
    idx = 1
    while idx < len(runs):
        curr_t = runs[idx].find(qn('w:t'))
        prev_t = runs[idx - 1].find(qn('w:t'))
        if (curr_t is not None and curr_t.text and prev_t is not None and prev_t.text
                and curr_t.text[0] in ALL_SPECIAL and prev_t.text[-1].isalpha()):
            moved = prev_t.text[-1]
            prev_t.text = prev_t.text[:-1]
            curr_t.text = moved + curr_t.text
        idx += 1


def process_run(run_elem, parent):
    text_elem = run_elem.find(qn('w:t'))
    if text_elem is None or not text_elem.text:
        return 0
    text = text_elem.text
    if not has_special(text):
        return 0
    rPr = copy_rPr(run_elem)
    segments = parse_segments(text)
    if not segments:
        return 0
    run_index = list(parent).index(run_elem)
    new_elems = []
    count = 0
    for seg in segments:
        if seg[0] == 'text':
            if seg[1]:
                new_elems.append(make_text_run(
                    seg[1],
                    etree.fromstring(etree.tostring(rPr)) if rPr is not None else None))
        elif seg[0] == 'sub':
            new_elems.append(make_omath_sub(seg[1], seg[2]))
            count += 1
        elif seg[0] == 'sup':
            new_elems.append(make_omath_sup(seg[1], seg[2]))
            count += 1
        elif seg[0] == 'sub_raw':
            new_elems.append(make_omath_sub(' ', seg[1]))
            count += 1
        elif seg[0] == 'sup_raw':
            new_elems.append(make_omath_sup(' ', seg[1]))
            count += 1
        elif seg[0] == 'frac':
            new_elems.append(make_omath_frac(seg[1], seg[2]))
            count += 1
    parent.remove(run_elem)
    for j, elem in enumerate(new_elems):
        parent.insert(run_index + j, elem)
    return count


# ============================================================
print("[1/2] Loading document...")
doc = Document(SRC)
print(f"  {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

print("[2/2] Converting to OMath...")
converted = 0

for i, p in enumerate(doc.paragraphs):
    style = p.style.name
    if 'toc' in style.lower():
        continue
    merge_split_subscripts(p._element)
    runs = list(p._element.findall(qn('w:r')))
    for run_elem in runs:
        t = run_elem.find(qn('w:t'))
        if t is not None and t.text and has_special(t.text):
            converted += process_run(run_elem, p._element)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                merge_split_subscripts(para._element)
                runs = list(para._element.findall(qn('w:r')))
                for run_elem in runs:
                    t = run_elem.find(qn('w:t'))
                    if t is not None and t.text and has_special(t.text):
                        converted += process_run(run_elem, para._element)

print(f"  OMath formulas created: {converted}")

remaining = 0
for p in doc.paragraphs:
    for run in p.runs:
        t = run._element.find(qn('w:t'))
        if t is not None and t.text:
            remaining += sum(1 for c in t.text if c in ALL_SPECIAL)
print(f"  Remaining unconverted: {remaining}")

print("Saving...")
doc.save(DST)
print(f"Saved to: {DST}")
